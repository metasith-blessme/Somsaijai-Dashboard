import os
import json
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_populated_sale_record(output_path, branch, data):
    doc = Document()
    
    # Header
    p = doc.add_paragraph()
    run = p.add_run(f"SomSaiJai Sales Record - {branch}")
    run.bold = True
    run.font.size = Pt(18)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Branch Circle
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run(f"({branch})")
    run.font.size = Pt(14)
    run.bold = True

    # Date/Day Box
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    cell = table.rows[0].cells[0]
    cell.width = Inches(2)
    p = cell.paragraphs[0]
    p.add_run(f"Date: {data['d']}\nDay: {data['day']}").font.size = Pt(12)

    doc.add_paragraph()

    # Financial Summary Section (Right Side)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.add_run(f"All ➔ {data['rev']} ฿").bold = True
    p.add_run(f"\n\nCash ➔ {data['cash']} ฿")
    p.add_run(f"\nice ➔ {data['exp']} ฿")
    p.add_run("\n-------------------------------").bold = True
    p.add_run(f"\nNet Cash ➔ {data['net']} ฿")
    p.add_run(f"\n\nScan ➔ {data['scan']} ฿").bold = True

    doc.add_paragraph()

    # Stock & Usage Section (Two Columns)
    table = doc.add_table(rows=1, cols=2)
    left_cell = table.rows[0].cells[0]
    right_cell = table.rows[0].cells[1]

    lp = left_cell.add_paragraph()
    lp.add_run("Cup Totals:").bold = True
    lp.add_run(f"\nOrange ➔ {data['or']}")
    lp.add_run(f"\nWater  ➔ {data['wm']}")
    lp.add_run(f"\nMango  ➔ {data['mg']}")
    lp.add_run(f"\nCoco   ➔ {data['co']}")
    lp.add_run(f"\nApple  ➔ {data['ap']}")
    if data.get('guava'): lp.add_run(f"\nGuava  ➔ {data['guava']}")
    if data.get('pineapple'): lp.add_run(f"\nPine   ➔ {data['pineapple']}")
    lp.add_run(f"\nTotal  ➔ {data['tot']}").bold = True

    lp.add_run("\n\nUsage:").bold = True
    lp.add_run(f"\nOrange ➔ {data['uo']} baskets")
    lp.add_run(f"\nWater  ➔ {data['uw']} pcs")
    lp.add_run(f"\nMango  ➔ {data['umg']} units")
    lp.add_run(f"\nApple  ➔ {data['uap']} units")

    rp = right_cell.add_paragraph()
    rp.add_run("Coco Details:").bold = True
    rp.add_run(f"\nMeat   ➔ {data.get('uco_meat', 0)}")
    rp.add_run(f"\nWater  ➔ {data.get('uco_water', 0)}")
    rp.add_run(f"\nRaw    ➔ {data.get('uco_raw', 0)}")
    rp.add_run(f"\nConden ➔ {data.get('uco_conden', 0)}")

    doc.save(output_path)

if __name__ == "__main__":
    docs_dir = "SomSaiJai_Documents"
    with open("3_Automation_Dashboard/data.json", 'r') as f:
        full_data = json.load(f)
    
    # Extract May 13 specifically
    b1_may13 = next(s for s in full_data['branches']['B1']['sales']['May26'] if s['d'] == "13/05/2026")
    b2_may13 = next(s for s in full_data['branches']['B2']['sales']['May26'] if s['d'] == "13/05/2026")
    
    create_populated_sale_record(f"{docs_dir}/B1_Sale_Record_May13.docx", "B1", b1_may13)
    create_populated_sale_record(f"{docs_dir}/B2_Sale_Record_May13.docx", "B2", b2_may13)
    print("Populated sale records for May 13 generated.")
