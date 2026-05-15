import os
import json
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_cell_border(cell, **kwargs):
    """
    Set cell border
    Usage: set_cell_border(cell, top={"sz": 12, "val": "single", "color": "#FF0000", "space": "0"}, ...)
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for side in ('top', 'left', 'bottom', 'right'):
        if side in kwargs:
            edge = OxmlElement(f'w:{side}')
            for key, val in kwargs[side].items():
                edge.set(qn(f'w:{key}'), str(val))
            tcPr.append(edge)

def create_sale_record_template(output_path, branch="B1"):
    doc = Document()
    
    # Header
    p = doc.add_paragraph()
    run = p.add_run(f"SomSaiJai Sales Record - {branch}")
    run.bold = True
    run.font.size = Pt(18)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Branch Circle (simulation)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run(f"({branch})")
    run.font.size = Pt(14)
    run.bold = True

    # Date/Day Box
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    cell = table.rows[0].cells[0]
    cell.width = Inches(2)
    p = cell.paragraphs[0]
    p.add_run("Date: ............................\nDay: .............................").font.size = Pt(12)

    doc.add_paragraph() # Spacer

    # Tally Section
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Item / Tally'
    hdr_cells[1].text = 'Total (฿)'
    hdr_cells[2].text = 'Cups'

    items = [
        "Cap(60) Cash", "Cap(60) Scan", "Cap(50) Cash", "Cap(50) Scan",
        "Mango(90) Cash", "Mango(90) Scan", "Coco(60) Cash", "Coco(60) Scan",
        "Apple(60) Cash", "Apple(60) Scan"
    ]

    for item in items:
        row_cells = table.add_row().cells
        row_cells[0].text = f"{item} - "
        row_cells[1].text = "(          )"
        row_cells[2].text = ""

    doc.add_paragraph()

    # Financial Summary Section (Right Side)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.add_run("All ➔ ...........................").bold = True
    p.add_run("\n\nCash ➔ ...........................")
    p.add_run("\nice ➔ ...........................")
    p.add_run("\n-------------------------------").bold = True
    p.add_run("\nNet Cash ➔ ...........................")
    p.add_run("\n\nScan ➔ ...........................")

    doc.add_paragraph()

    # Stock & Usage Section (Two Columns)
    table = doc.add_table(rows=1, cols=2)
    left_cell = table.rows[0].cells[0]
    right_cell = table.rows[0].cells[1]

    lp = left_cell.add_paragraph()
    lp.add_run("Cup Totals:").bold = True
    lp.add_run("\nOrange ➔ (      )")
    lp.add_run("\nWater  ➔ (      )")
    lp.add_run("\nMango  ➔ (      )")
    lp.add_run("\nCoco   ➔ (      )")
    lp.add_run("\nApple  ➔ (      )")
    lp.add_run("\nTotal  ➔ (      )").bold = True

    lp.add_run("\n\nUsage:").bold = True
    lp.add_run("\nOrange ➔ (      ) baskets")
    lp.add_run("\nWater  ➔ (      ) pcs")
    lp.add_run("\nMango  ➔ (      ) units")
    lp.add_run("\nApple  ➔ (      ) units")

    rp = right_cell.add_paragraph()
    rp.add_run("Coco Details:").bold = True
    rp.add_run("\nMeat   ➔ (      )")
    rp.add_run("\nWater  ➔ (      )")
    rp.add_run("\nRaw    ➔ (      )")
    rp.add_run("\nConden ➔ (      )")

    doc.save(output_path)

def create_stock_report(output_path, ledger_path):
    doc = Document()
    doc.add_heading('SomSaiJai Stock Inventory Report', 0)

    if os.path.exists(ledger_path):
        with open(ledger_path, 'r') as f:
            ledger = json.load(f)
        
        # Physical Checks
        doc.add_heading('Latest Physical Checks', level=1)
        table = doc.add_table(rows=1, cols=4)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Date'
        hdr_cells[1].text = 'Item'
        hdr_cells[2].text = 'Quantity'
        hdr_cells[3].text = 'Unit'

        for check in ledger.get('physical_checks', [])[-10:]:
            row_cells = table.add_row().cells
            row_cells[0].text = check.get('date', '')
            row_cells[1].text = check.get('item', '')
            row_cells[2].text = str(check.get('qty', ''))
            row_cells[3].text = check.get('unit', '')

        # Purchases
        doc.add_heading('Recent Purchases', level=1)
        table = doc.add_table(rows=1, cols=3)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Date'
        hdr_cells[1].text = 'Item'
        hdr_cells[2].text = 'Quantity'

        for purchase in ledger.get('purchases', [])[-5:]:
            row_cells = table.add_row().cells
            row_cells[0].text = purchase.get('date', '')
            row_cells[1].text = purchase.get('item', '')
            row_cells[2].text = str(purchase.get('qty', ''))
    else:
        doc.add_paragraph("Stock ledger file not found.")

    doc.save(output_path)

def create_attendance_template(output_path):
    doc = Document()
    doc.add_heading('Employee Attendance Record', 0)
    
    p = doc.add_paragraph()
    p.add_run("Month: ............................ Year: ............................").bold = True
    
    table = doc.add_table(rows=1, cols=7)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Date'
    hdr_cells[1].text = 'Name'
    hdr_cells[2].text = 'Time In'
    hdr_cells[3].text = 'Time Out'
    hdr_cells[4].text = 'Break'
    hdr_cells[5].text = 'Total Hours'
    hdr_cells[6].text = 'Signature'

    # Add 31 rows for a full month
    for i in range(1, 32):
        row_cells = table.add_row().cells
        row_cells[0].text = str(i)

    doc.add_paragraph("\n\n")
    p = doc.add_paragraph()
    p.add_run("Manager Signature: ...........................................").italic = True
    
    doc.save(output_path)

if __name__ == "__main__":
    docs_dir = "SomSaiJai_Documents"
    create_sale_record_template(f"{docs_dir}/Daily_Sale_Record_Template.docx")
    create_stock_report(f"{docs_dir}/Stock_Inventory_Report.docx", "3_Automation_Dashboard/stock_ledger.json")
    create_attendance_template(f"{docs_dir}/Employee_Attendance_Template.docx")
    print("Documents generated successfully.")
